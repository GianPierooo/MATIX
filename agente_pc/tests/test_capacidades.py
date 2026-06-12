"""Capacidades TIPADAS (librería confiable por tarea). Deterministas, validadas,
con rieles: rutas por denylist, sin shell, resultado claro. Sin abrir nada real
(el abridor/capturador se inyectan en el Contexto)."""
from __future__ import annotations

import os
from pathlib import Path

from agente_pc import capacidades as cap
from agente_pc.acciones import crear_registro
from agente_pc.registro import Contexto, NivelRiesgo

_HOME = Path.home()


def _ctx(**kw) -> Contexto:
    return Contexto(allowlist=[_HOME], **kw)


# ── Registro: las capacidades quedan registradas con su nivel ────────────────


def test_capacidades_en_el_registro():
    # Todo lo REVERSIBLE es SEGURA: directo, sin fricción de confirmación.
    reg = crear_registro()
    assert reg.get("abrir_carpeta").nivel is NivelRiesgo.SEGURA
    assert reg.get("tomar_captura").nivel is NivelRiesgo.SEGURA
    assert reg.get("crear_documento_word").nivel is NivelRiesgo.SEGURA
    assert reg.get("reproducir_spotify").nivel is NivelRiesgo.SEGURA
    assert reg.get("verificar_spotify").nivel is NivelRiesgo.SEGURA
    assert reg.get("abrir_web").nivel is NivelRiesgo.SEGURA


# ── abrir_web ────────────────────────────────────────────────────────────────


def test_abrir_web_http_ok():
    ctx = _ctx()
    abierto = {}
    ctx.abridor = lambda u: abierto.update(url=u) or {"ok": True}
    r = cap._abrir_web({"url": "https://github.com/x"}, ctx)
    assert r["ok"] and r["tipo"] == "web_abierta"
    assert abierto["url"] == "https://github.com/x"


def test_abrir_web_sin_esquema_asume_https():
    ctx = _ctx()
    abierto = {}
    ctx.abridor = lambda u: abierto.update(url=u) or {"ok": True}
    r = cap._abrir_web({"url": "youtube.com"}, ctx)
    assert r["ok"] and abierto["url"] == "https://youtube.com"


def test_abrir_web_rechaza_no_http():
    ctx = _ctx()
    ctx.abridor = lambda u: (_ for _ in ()).throw(AssertionError("no debe abrir"))
    for mala in ("file:///C:/Users/gianp/.ssh/id_rsa", "javascript:alert(1)", "data:text/html,x"):
        r = cap._abrir_web({"url": mala}, ctx)
        assert not r["ok"] and r["tipo"] == "validacion", mala


def test_abrir_web_sin_url():
    r = cap._abrir_web({}, _ctx())
    assert not r["ok"] and r["tipo"] == "validacion"


# ── abrir_carpeta ────────────────────────────────────────────────────────────


def test_abrir_carpeta_permitida(tmp_path, monkeypatch):
    # tmp_path cae bajo AppData (denylisted), así que usamos una carpeta del home.
    d = _HOME / "Documents"
    if not d.is_dir():
        return  # entorno sin Documents; el resto de tests cubren la lógica
    ctx = _ctx()
    ctx.abridor = lambda ruta: {"ok": True}
    r = cap._abrir_carpeta({"ruta": str(d)}, ctx)
    assert r["ok"] and r["tipo"] == "abierto" and r["es_carpeta"] is True


def test_abrir_carpeta_denylist():
    # `.ssh` es un componente prohibido cross-platform (no como C:\Windows, que
    # solo aplica en Windows): la denylist lo bloquea aunque esté en el home.
    ctx = _ctx()
    ctx.abridor = lambda ruta: {"ok": True}
    r = cap._abrir_carpeta({"ruta": str(_HOME / ".ssh")}, ctx)
    assert not r["ok"] and r["tipo"] == "rechazada"


def test_abrir_carpeta_sin_ruta():
    r = cap._abrir_carpeta({}, _ctx())
    assert not r["ok"] and r["tipo"] == "validacion"


# ── crear_documento_word ─────────────────────────────────────────────────────


def test_crear_word_genera_docx_real():
    carpeta = _HOME / "Documents"
    if not carpeta.is_dir():
        return
    ctx = _ctx()
    r = cap._crear_documento_word({
        "titulo": "Reporte test",
        "parrafos": ["Uno.", "Dos."],
        "tablas": [{"encabezados": ["A", "B"], "filas": [["1", "2"], ["3", "4"]]}],
        "carpeta": str(carpeta),
        "nombre": "_matix_test_doc",
    }, ctx)
    assert r["ok"] and r["tipo"] == "documento_creado"
    try:
        from docx import Document
        d = Document(r["ruta"])
        assert any("Reporte test" in p.text for p in d.paragraphs)
        assert len(d.tables) == 1 and len(d.tables[0].rows) == 3  # encabezado + 2
    finally:
        if os.path.isfile(r.get("ruta", "")):
            os.remove(r["ruta"])


def test_crear_word_vacio_rechaza():
    r = cap._crear_documento_word({}, _ctx())
    assert not r["ok"] and r["tipo"] == "validacion"


def test_crear_word_carpeta_denylisted():
    ctx = _ctx()
    r = cap._crear_documento_word(
        {"titulo": "x", "carpeta": str(_HOME / "AppData" / "Local")}, ctx
    )
    assert not r["ok"] and r["tipo"] == "rechazada"


# ── reproducir_spotify ───────────────────────────────────────────────────────


def test_spotify_consulta_arma_uri_busqueda():
    ctx = _ctx()
    capturado = {}
    ctx.abridor = lambda uri: capturado.update(uri=uri) or {"ok": True}
    ctx.verificador_spotify = lambda *_a, **_k: (_ for _ in ()).throw(AssertionError("no medir"))
    r = cap._reproducir_spotify({"consulta": "Michael Jackson"}, ctx)
    assert r["ok"] and r["tipo"] == "spotify_abierto"
    assert capturado["uri"].startswith("spotify:search:")
    assert "Michael" in capturado["uri"]
    # Una búsqueda nunca reproduce sola: lo dice honesto, sin gastar la espera.
    assert r["sonando"] is False


def test_spotify_uri_track_verifica_si_suena():
    ctx = _ctx()
    ctx.abridor = lambda uri: {"ok": True}
    ctx.verificador_spotify = lambda espera: {"sonando": True, "titulo": "MJ - Billie Jean"}
    r = cap._reproducir_spotify({"uri": "spotify:track:abc"}, ctx)
    assert r["ok"] and r["uri"] == "spotify:track:abc"
    assert r["sonando"] is True and r["reproduciendo"] == "MJ - Billie Jean"


def test_spotify_uri_track_honesto_si_no_suena():
    ctx = _ctx()
    ctx.abridor = lambda uri: {"ok": True}
    ctx.verificador_spotify = lambda espera: {"sonando": False, "titulo": None}
    r = cap._reproducir_spotify({"uri": "spotify:track:abc"}, ctx)
    assert r["ok"] and r["sonando"] is False  # abrió, pero NO miente con «suena»


def test_verificar_spotify_solo_mide():
    ctx = _ctx()
    ctx.verificador_spotify = lambda espera: {"sonando": True, "titulo": "X - Y"}
    r = cap._verificar_spotify({"espera_s": 2}, ctx)
    assert r["ok"] and r["tipo"] == "spotify_verificado"
    assert r["sonando"] is True and r["reproduciendo"] == "X - Y"


def test_spotify_uri_no_protocolo_rechaza():
    ctx = _ctx()
    ctx.abridor = lambda uri: {"ok": True}
    r = cap._reproducir_spotify({"uri": "http://evil.com"}, ctx)
    assert not r["ok"] and r["tipo"] == "validacion"


def test_spotify_sin_datos_rechaza():
    r = cap._reproducir_spotify({}, _ctx())
    assert not r["ok"] and r["tipo"] == "validacion"


# ── tomar_captura (capturador inyectado) ─────────────────────────────────────


def test_tomar_captura_guarda_y_devuelve_ruta():
    carpeta = _HOME / "Pictures"
    ctx = _ctx()
    guardado = {}
    ctx.capturador_archivo = lambda ruta: guardado.update(ruta=ruta) or {"ok": True, "ancho": 1920, "alto": 1080}
    r = cap._tomar_captura({}, ctx)
    assert r["ok"] and r["tipo"] == "captura_tomada"
    assert r["ruta"].endswith(".png") and "Matix" in r["ruta"]
    assert r["ancho"] == 1920
