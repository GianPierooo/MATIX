"""Capacidades TIPADAS de alto nivel (Capa 6 · librería de capacidades).

Lección de arquitectura: la fiabilidad NO viene de controlar la pantalla a
ciegas, sino de una herramienta CONFIABLE por tarea — determinista, validada,
con resultado claro. El control de pantalla (`pantalla.py`) queda como ÚLTIMO
RECURSO cuando no hay una capacidad nativa.

Este módulo agrega el primer lote:
  - abrir_carpeta — abre una ruta permitida en el Explorador (sin shell).
  - tomar_captura — screenshot a PNG en una carpeta del usuario, devuelve ruta.
  - crear_documento_word — .docx real con python-docx (título, párrafos, tablas).
  - reproducir_spotify — abre Spotify por URI `spotify:` (determinista, no clics).

SEGURIDAD: todo se valida en el BORDE (aquí), nunca confiando en el modelo. Las
rutas pasan por `seguridad.ruta_permitida` (denylist gana sobre allowlist); no se
ejecuta shell; las consecuentes exigen `confirmado=true` (lo aplica el registry).
"""
from __future__ import annotations

import logging
import os
import re
import subprocess
import sys
from datetime import datetime, timedelta, timezone
from typing import Any
from urllib.parse import quote

from .registro import AccionDef, Contexto, NivelRiesgo, Param
from .seguridad import ruta_permitida

log = logging.getLogger("matix.agente")

_LIMA = timezone(timedelta(hours=-5))  # Perú: UTC-5 fijo, sin tzdata

# Carpetas comunes por nombre amistoso (mismo criterio que acciones.py).
_CARPETAS = {
    "documentos": "~/Documents", "documents": "~/Documents",
    "escritorio": "~/Desktop", "desktop": "~/Desktop",
    "descargas": "~/Downloads", "downloads": "~/Downloads",
    "imagenes": "~/Pictures", "imágenes": "~/Pictures", "pictures": "~/Pictures",
    "videos": "~/Videos", "música": "~/Music", "musica": "~/Music",
}


def _err(tipo: str, mensaje: str, **extra: Any) -> dict[str, Any]:
    return {"ok": False, "tipo": tipo, "mensaje": mensaje, **extra}


def _resolver_nombre(ruta: str) -> str:
    bruto = (ruta or "").strip()
    clave = bruto.lower().strip("/\\")
    if not os.path.isabs(bruto) and clave in _CARPETAS:
        return _CARPETAS[clave]
    return bruto


def _real(ruta: str) -> str:
    return os.path.realpath(os.path.expanduser(_resolver_nombre(str(ruta))))


# ── abrir_carpeta ────────────────────────────────────────────────────────────


def _abrir_en_explorador(ruta_real: str) -> dict[str, Any]:
    """Abre `ruta_real` en el Explorador SIN shell. Inyectable vía override en
    los tests (monkeypatch). Real: os.startfile en Windows, xdg-open/open fuera."""
    try:
        if os.name == "nt":
            os.startfile(ruta_real)  # ShellExecute "open" — no es shell de comandos
        elif sys.platform == "darwin":
            subprocess.Popen(["open", ruta_real], shell=False)
        else:
            subprocess.Popen(["xdg-open", ruta_real], shell=False)
        return {"ok": True}
    except OSError as e:
        return {"ok": False, "error": type(e).__name__}


def _abrir_carpeta(args: dict[str, Any], ctx: Contexto) -> dict[str, Any]:
    ruta = _resolver_nombre((args or {}).get("ruta"))
    if not ruta:
        return _err("validacion", "necesito la ruta a abrir")
    if not ruta_permitida(ruta, ctx.allowlist).permitida:
        return _err("rechazada", "esa ruta no está permitida (denylist o fuera de tu perfil)")
    real = _real(ruta)
    # Abre carpetas (en el Explorador) o ARCHIVOS (en su app por defecto: un
    # .docx → Word, un .pdf → el lector…). Ambos vía ShellExecute, sin shell.
    es_dir = os.path.isdir(real)
    es_archivo = os.path.isfile(real)
    if not es_dir and not es_archivo:
        return _err("no_existe", "no encontré esa carpeta o archivo")
    abridor = getattr(ctx, "abridor", None) or _abrir_en_explorador
    res = abridor(real)
    if not res.get("ok"):
        return _err("error_abrir", "no pude abrir eso")
    return {"ok": True, "tipo": "abierto", "ruta": real, "es_carpeta": es_dir}


# ── tomar_captura (screenshot a PNG) ─────────────────────────────────────────


def _tomar_captura(args: dict[str, Any], ctx: Contexto) -> dict[str, Any]:
    # Carpeta destino: por defecto ~/Pictures/Matix (dentro del perfil; la
    # allowlist/denylist la cubre). Se valida igual.
    carpeta_in = (args or {}).get("carpeta") or "~/Pictures"
    carpeta = _real(carpeta_in)
    if not ruta_permitida(carpeta, ctx.allowlist).permitida:
        return _err("rechazada", "la carpeta destino de la captura no está permitida")
    destino_dir = os.path.join(carpeta, "Matix")
    nombre = datetime.now(_LIMA).strftime("captura_%Y-%m-%d_%H-%M-%S.png")
    ruta_png = os.path.join(destino_dir, nombre)
    capturador = getattr(ctx, "capturador_archivo", None)
    try:
        if capturador is not None:
            res = capturador(ruta_png)
        else:
            from .pantalla import capturar_a_archivo
            res = capturar_a_archivo(ruta_png)
    except Exception as e:  # noqa: BLE001 — nunca tumbar el proceso
        log.exception("tomar_captura falló")
        return _err("error_captura", f"no pude tomar la captura ({type(e).__name__})")
    if not res.get("ok"):
        return _err(res.get("tipo", "error_captura"), res.get("mensaje", "no pude capturar"))
    return {
        "ok": True, "tipo": "captura_tomada",
        "ruta": ruta_png, "ancho": res.get("ancho"), "alto": res.get("alto"),
    }


# ── crear_documento_word (python-docx) ───────────────────────────────────────


def _nombre_archivo_seguro(nombre: str, defecto: str) -> str:
    base = (nombre or "").strip() or defecto
    base = os.path.splitext(base)[0]  # sin extensión que ponga el usuario
    # Solo caracteres seguros de nombre; nada de separadores ni '..'.
    base = re.sub(r"[^\w\sáéíóúñÁÉÍÓÚÑ.-]", "", base, flags=re.UNICODE).strip()
    base = base.replace("..", "").strip(". ") or defecto
    return base[:120] + ".docx"


def _crear_documento_word(args: dict[str, Any], ctx: Contexto) -> dict[str, Any]:
    args = args or {}
    titulo = str(args.get("titulo") or "").strip()
    parrafos = args.get("parrafos") or []
    tablas = args.get("tablas") or []
    if not titulo and not parrafos and not tablas:
        return _err("validacion", "dame al menos un título, párrafos o una tabla para el documento")
    if not isinstance(parrafos, list) or not isinstance(tablas, list):
        return _err("validacion", "«parrafos» y «tablas» deben ser listas")

    # Carpeta destino: ~/Documents por defecto, validada.
    carpeta = _real((args.get("carpeta") or "~/Documents"))
    if not ruta_permitida(carpeta, ctx.allowlist).permitida:
        return _err("rechazada", "la carpeta destino del documento no está permitida")
    if not os.path.isdir(carpeta):
        return _err("no_existe", "la carpeta destino no existe")
    nombre = _nombre_archivo_seguro(args.get("nombre") or titulo, "documento_matix")
    ruta = os.path.join(carpeta, nombre)
    if os.path.exists(ruta):
        # No sobreescribir: agrega un sufijo de hora.
        sufijo = datetime.now(_LIMA).strftime("_%H-%M-%S")
        ruta = os.path.join(carpeta, nombre[:-5] + sufijo + ".docx")

    try:
        from docx import Document  # type: ignore
    except Exception:  # noqa: BLE001
        return _err(
            "sin_docx",
            "falta python-docx en la PC para crear Word. Instálalo: "
            "cd agente_pc && uv sync",
        )

    try:
        doc = Document()
        if titulo:
            doc.add_heading(titulo, level=0)
        for p in parrafos:
            doc.add_paragraph(str(p))
        for tabla in tablas:
            t = tabla if isinstance(tabla, dict) else {}
            sub = str(t.get("titulo") or "").strip()
            if sub:
                doc.add_heading(sub, level=2)
            encabezados = [str(h) for h in (t.get("encabezados") or [])]
            filas = [[str(c) for c in fila] for fila in (t.get("filas") or []) if isinstance(fila, list)]
            ncols = max([len(encabezados)] + [len(f) for f in filas] or [0])
            if ncols == 0:
                continue
            tab = doc.add_table(rows=0, cols=ncols)
            tab.style = "Light Grid Accent 1"
            if encabezados:
                celdas = tab.add_row().cells
                for i, h in enumerate(encabezados[:ncols]):
                    celdas[i].text = h
            for fila in filas:
                celdas = tab.add_row().cells
                for i, c in enumerate(fila[:ncols]):
                    celdas[i].text = c
        doc.save(ruta)
    except Exception as e:  # noqa: BLE001
        log.exception("crear_documento_word falló")
        return _err("error_docx", f"no pude crear el documento ({type(e).__name__})")

    return {"ok": True, "tipo": "documento_creado", "ruta": ruta, "nombre": os.path.basename(ruta)}


# ── reproducir_spotify (URI spotify:) ────────────────────────────────────────


def _abrir_uri(uri: str) -> dict[str, Any]:
    """Abre un URI de protocolo (spotify:…) con el handler del SO. Sin shell."""
    try:
        if os.name == "nt":
            os.startfile(uri)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", uri], shell=False)
        else:
            subprocess.Popen(["xdg-open", uri], shell=False)
        return {"ok": True}
    except OSError as e:
        return {"ok": False, "error": type(e).__name__}


# URIs que apuntan a algo REPRODUCIBLE (vs. spotify:search:, que solo navega).
_SPOTIFY_REPRODUCIBLES = {"track", "album", "playlist", "episode", "show"}


def _verificador(ctx: Contexto):
    inyectado = getattr(ctx, "verificador_spotify", None)
    if inyectado is not None:
        return inyectado
    from . import audio
    return audio.verificar_sonando


def _reproducir_spotify(args: dict[str, Any], ctx: Contexto) -> dict[str, Any]:
    args = args or {}
    uri = (args.get("uri") or "").strip()
    consulta = (args.get("consulta") or "").strip()
    if uri:
        if not uri.startswith("spotify:"):
            return _err("validacion", "el uri debe empezar con «spotify:»")
        destino = uri
        humano = uri
    elif consulta:
        # Determinista: abre Spotify en la búsqueda exacta. No clica a ciegas.
        destino = "spotify:search:" + quote(consulta)
        humano = f"búsqueda «{consulta}»"
    else:
        return _err("validacion", "dime qué reproducir (una canción/artista) o un uri de Spotify")
    abridor = getattr(ctx, "abridor", None) or _abrir_uri
    res = abridor(destino)
    if not res.get("ok"):
        return _err("error_spotify", "no pude abrir Spotify; ¿está instalado?")
    salida: dict[str, Any] = {"ok": True, "tipo": "spotify_abierto", "uri": destino, "detalle": humano}
    # HONESTIDAD: si el destino es reproducible, MEDIMOS si de verdad suena
    # (peak de audio + título de ventana). Abrir spotify:track:… NAVEGA pero el
    # cliente no auto-reproduce; jamás reportamos «sonando» sin medirlo.
    partes = destino.split(":")
    if len(partes) >= 2 and partes[1] in _SPOTIFY_REPRODUCIBLES:
        v = _verificador(ctx)(float(args.get("espera_s") or 8.0)) or {}
        salida["sonando"] = v.get("sonando")
        salida["reproduciendo"] = v.get("titulo")
    else:
        salida["sonando"] = False  # una búsqueda nunca reproduce sola
        salida["reproduciendo"] = None
    return salida


def _verificar_spotify(args: dict[str, Any], ctx: Contexto) -> dict[str, Any]:
    """Solo MIDE si Spotify está sonando (peak + título). No abre nada. Lo usa
    el cerebro para re-verificar tras ordenar play por la Web API."""
    v = _verificador(ctx)(float((args or {}).get("espera_s") or 6.0)) or {}
    return {
        "ok": True, "tipo": "spotify_verificado",
        "sonando": v.get("sonando"), "reproduciendo": v.get("titulo"),
    }


# ── Definiciones ──────────────────────────────────────────────────────────────

# NIVELES: lo REVERSIBLE (abrir, reproducir, capturar, crear un doc nuevo que
# nunca sobreescribe) es SEGURA — se ejecuta directo, sin fricción de
# confirmación. La confirmación queda para lo irreversible/grave (borrar,
# sobrescribir, cerrar apps con trabajo a medias), no para abrir una carpeta.
DEFS_CAPACIDADES: list[AccionDef] = [
    AccionDef(
        "abrir_carpeta",
        "Abre una carpeta del usuario en el Explorador (ruta permitida). Si le "
        "pasas un archivo, abre su carpeta. Determinista, sin tocar la pantalla.",
        (Param("ruta", str, requerido=True),),
        NivelRiesgo.SEGURA,
        _abrir_carpeta,
    ),
    AccionDef(
        "tomar_captura",
        "Toma una captura de pantalla y la guarda como PNG en ~/Pictures/Matix; "
        "devuelve la ruta. Solo lectura de la pantalla, determinista.",
        (Param("carpeta", str, requerido=False),),
        NivelRiesgo.SEGURA,
        _tomar_captura,
    ),
    AccionDef(
        "crear_documento_word",
        "Crea un .docx REAL con python-docx: título, párrafos y tablas con los "
        "datos dados. NO maneja la GUI de Word. Lo guarda y devuelve la ruta.",
        (
            Param("titulo", str, requerido=False),
            Param("parrafos", list, requerido=False),
            Param("tablas", list, requerido=False),
            Param("nombre", str, requerido=False),
            Param("carpeta", str, requerido=False),
        ),
        NivelRiesgo.SEGURA,
        _crear_documento_word,
    ),
    AccionDef(
        "reproducir_spotify",
        "Abre Spotify por URI (spotify:search:… o un uri dado) y VERIFICA si "
        "suena (peak de audio + título de ventana). Determinista, sin clics.",
        # espera_s llega sin declarar (el handler la coerciona con float()).
        (Param("consulta", str, requerido=False), Param("uri", str, requerido=False)),
        NivelRiesgo.SEGURA,
        _reproducir_spotify,
    ),
    AccionDef(
        "verificar_spotify",
        "Mide si Spotify está sonando en este momento (peak + título de "
        "ventana). Solo lectura; no abre ni toca nada.",
        (),
        NivelRiesgo.SEGURA,
        _verificar_spotify,
    ),
]
